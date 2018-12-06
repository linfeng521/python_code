# coding:utf-8
import xlrd
import docx
class Student(object):
    def __init__(self,id,name):
        self.id = id
        self.name = name

    def __repr__(self):
        return 'student:'+str(self.id)+'\t'+self.name

students = []

def read_excel():
    # 打开文件
    workbook = xlrd.open_workbook(r'1.xls')
    # 根据sheet索引或者名称获取sheet内容
    sheet2 = workbook.sheet_by_index(0)  # sheet索引从0开始
    for i in range(2,32,1):
        list = sheet2.row_values(i)
        student = Student(int(list[0]),list[1])
        students.append(student)

def write_docx(student,i):
    from docx.oxml.ns import qn
    from docx.shared import Pt
    document = docx.Document('%i.docx'%(i))
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    docText2 = [paragraph.text for paragraph in document.paragraphs]
    # id
    id = document.paragraphs[9]
    run1 = id.add_run(str(student.id))
    font = run1.font
    font.name = '宋体'
    font.size = Pt(16)
    font.underline = True
    # run1= id.add_run('专业班级：')
    # font = run1.font
    # font.name = '宋体'
    # font.size = Pt(16)
    # run1 = id.add_run('')
    # font = run1.font
    # font.name = '宋体'
    # font.size = Pt(16)

    # name
    p = document.paragraphs[10].clear()
    run1 = p.add_run('学生姓名：')
    font = run1.font
    font.name = '宋体'
    font.size = Pt(16)
    run1 = p.add_run('\t\t\t\t'+student.name+'\t\t\t\t\t')
    font = run1.font
    font.name = '宋体'
    font.size = Pt(16)
    font.underline = True
    filename = str(student.id)+'-'+student.name+'.docx'
    document.save(filename)
    # a = 0
    # for i in docText2:
    #     print('第'+str(a)+'位置是'+i)
    #     a = a+1
        #print(docText2.index('专业班级： 机械1607班   学号： 2016000'))
   # p = document.paragraphs[22].clear()
    # run1 = p.add_run('XXXX:        ')
    # font = run1.font
    # font.size = Pt(12)
    # font.bold = True
    # run2 = p.add_run(str_to_write)
    # run2.underline = True
    # run2.size = Pt(12)
if __name__ == '__main__':
    s = Student(1,'linfeng')
    print(s)
    read_excel()
    li = (1,2,3,4,5)
    li2 = ()
	#列表融合
    for i in range(1,7):
        li2=li2+li
    for i in range(30):
        write_docx(students[i],li2[i])
        print(students[i].name+'has been writen successfully')

    # for student,i in students,li2:
    #     #write_docx(student)
    #     print(student+i)
